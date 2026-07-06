from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "contest"
OUT_PATH = OUT_DIR / "SupplyMapAI_공모전_기획서_작성안.docx"
LOGO_PATH = ROOT / "public" / "portmind-logo.png"


BLUE = "1F5FBF"
INK = "17212B"
MUTED = "66727F"
TEAL = "0F766E"
AMBER = "B7791F"
LIGHT_BLUE = "EEF4FB"
LIGHT_GRAY = "F4F6F8"
LINE = "D8E0E6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Malgun Gothic"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), "Malgun Gothic")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = LINE) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def style_run(run, size: int = 10, color: str = INK, bold: bool = False) -> None:
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_para(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        style_run(r2)
    else:
        run = p.add_run(text)
        style_run(run)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        style_run(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        style_run(run)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=16 if level == 1 else 13 if level == 2 else 11, color=BLUE if level <= 2 else INK, bold=True)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=fill)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    style_run(r, size=10, color=BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    style_run(r2, size=9, color=INK)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None, font_size: int = 8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths_cm:
        set_column_widths(table, widths_cm)

    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        set_cell_shading(header.cells[idx], LIGHT_GRAY)
        set_cell_text(header.cells[idx], text, bold=True, size=font_size)

    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            set_cell_text(row.cells[idx], text, size=font_size)
    doc.add_paragraph()
    return table


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style.font.size = Pt(10)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_cover(doc: Document) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.9))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("SupplyMap AI")
    style_run(r, size=28, color=INK, bold=True)

    p = doc.add_paragraph()
    r = p.add_run("산업통상부 공공데이터 활용 아이디어 공모전 기획서 작성안")
    style_run(r, size=14, color=BLUE, bold=True)

    add_callout(
        doc,
        "서비스 정의",
        "공공데이터 기반 한국·중국 공장 비교 및 RAG 무역 코파일럿. 사용자가 제품명, HS코드, 수입 희망 국가를 입력하면 국내 공장 후보, 중국 베타 공장 후보, 인증·통관·리콜·국가위험 근거를 결합해 AI 답변과 Evidence Drawer로 제공합니다.",
    )

    meta_rows = [
        ["제품/서비스명", "SupplyMap AI"],
        ["회사/브랜드", "PortMind"],
        ["현재 개발 상태", "Next.js App Router + Prisma + SQLite 기반 작동형 웹 MVP"],
        ["로컬 확인 URL", "http://127.0.0.1:3001/supplymap"],
        ["제출용 서비스 URL", "Vercel 배포 후 https://<project-name>.vercel.app/supplymap 입력"],
        ["GitHub URL", "GitHub 업로드 후 https://github.com/<user>/<repo> 입력"],
        ["작성 기준일", "2026년 7월 6일"],
    ]
    add_table(doc, ["항목", "기재 내용"], meta_rows, widths_cm=[4.0, 13.0], font_size=9)

    doc.add_page_break()


def add_application_form_text(doc: Document) -> None:
    add_heading(doc, "A. 참가 신청서 첫 페이지 기재안", 1)
    add_heading(doc, "제품서비스명", 2)
    add_para(doc, "SupplyMap AI")

    add_heading(doc, "세부내용 3줄 요약", 2)
    add_bullets(
        doc,
        [
            "산업통상부 및 산하 공공기관 공공데이터와 관세청·식약처 등 보조 데이터를 통합해 제품별 국내 공장 후보와 중국 베타 공장 후보를 탐색·비교하는 플랫폼입니다.",
            "RAG 기반 AI 무역 챗봇이 제품명, HS코드, 수입국, 공급 희망 지역을 구조화하고 인증·통관·리콜·국가위험 근거를 출처와 함께 요약합니다.",
            "중소기업·소상공인이 수입·조달 의사결정 전에 공장 후보, 지도 분포, 리스크, 확인 서류, 다음 액션을 한 화면에서 확인하도록 지원합니다.",
        ],
    )

    add_heading(doc, "산업통상부 활용 데이터", 2)
    add_bullets(
        doc,
        [
            "산업통상부 산하기관(한국산업단지공단) 공장등록생산정보조회서비스",
            "산업통상부 산하기관(한국산업단지공단) 산업동향조사 통계 조회 서비스",
            "산업통상부 산하기관(대한무역투자진흥공사) 무역투자 질의응답 데이터셋",
            "산업통상부 산하기관(대한무역투자진흥공사) 해외시장뉴스",
            "산업통상부 산하기관(대한무역투자진흥공사) 국가정보",
            "산업통상부 산하기관(한국무역보험공사) 국가별 업종별 위험지수 및 무역지표",
            "산업통상부 산하기관(국가기술표준원·Safety Korea) 제품안전 인증 및 리콜 정보",
        ],
    )

    add_heading(doc, "기타 활용 데이터", 2)
    add_bullets(
        doc,
        [
            "관세청: HS부호 및 세관장확인대상물품, 품목별 국가별 수출입실적",
            "관세청 UNI-PASS: HS부호검색, 세관장확인대상 물품 조회, 관세율 기본 조회, 관세환율 정보, HS CODE 내비게이션",
            "식품의약품안전처: 수입식품 해외제조업소 정보, 화장품 원료성분정보, 화장품 사용제한 원료정보, 의료기기 품목정보·품목허가정보",
            "사용자 업로드 데이터: 해외 제조업소·해외작업장·수출업소 엑셀 스냅샷, 공급업체 목록, 견적서, 카탈로그",
            "민간·지도 데이터: Kakao Map/Local, Google Maps Platform, AMap, 민간 해외 공장 후보 데이터",
            "LLM API: ChatGPT API 또는 DeepSeek API를 RAG 답변 생성 엔진으로 사용하되, API 키는 서버 환경변수로만 관리합니다.",
        ],
    )


def add_planning_sections(doc: Document) -> None:
    add_heading(doc, "B. 첨부3 기획서 작성안", 1)
    add_heading(doc, "제품/서비스명", 2)
    add_para(doc, "SupplyMap AI - 공공데이터 기반 한국·중국 공장 비교 및 RAG 무역 코파일럿")

    add_heading(doc, "제품/서비스의 간략 개요 3줄 이내", 2)
    add_bullets(
        doc,
        [
            "SupplyMap AI는 산업통상부·산하기관 공공데이터를 중심으로 국내 공장 후보와 중국 베타 공장 후보를 품목 기준으로 비교하는 공급망 탐색 플랫폼입니다.",
            "제품명과 HS코드 후보를 입력하면 공장 지도, 후보 비교 점수, 인증·통관·리콜·국가위험을 통합 분석하고 RAG 기반 AI 챗봇이 근거 있는 답변을 제공합니다.",
            "모든 후보·리스크·AI 답변에 출처 배지와 Evidence Drawer를 붙여 중소기업이 수입·조달 전 확인해야 할 데이터를 빠르게 파악하도록 돕습니다.",
        ],
    )

    add_heading(doc, "1. 제품 및 서비스 핵심내용(구체성, 우수성)", 2)
    add_para(
        doc,
        "SupplyMap AI는 기존의 단순 공장 검색이나 지도 표시를 넘어, 공공데이터를 의사결정 가능한 구조로 재가공하는 공급망 인텔리전스 서비스입니다. 사용자는 제품명, HS코드 후보, 수입 희망 국가, 국내 희망 권역을 입력하고, 시스템은 제품 키워드와 품목군을 추출한 뒤 국내 공장 후보, 중국 베타 공장 후보, 인증·통관·리콜·국가위험을 동시에 비교합니다.",
    )
    add_bullets(
        doc,
        [
            "국내 공장 히트맵: 한국산업단지공단 공장등록생산정보 기반 Supplier 3,322개를 Kakao Map 위에 표시합니다. 화장품, 식품 포장, 전기·전자, 의료·바이오, 드론·부품 등 카테고리 버튼으로 관련 공장만 필터링할 수 있습니다.",
            "중국 공장 베타 히트맵: 기존 Factory/Product DB의 중국 공장 42,389개 중 좌표 보유 37,723개를 활용합니다. 기구·용기·포장, 가공식품, 농산물, 식품첨가물, 수산물, 건강기능식품 등 품목군별 분포를 보여주되 좌표 추정 여부를 명확히 표시합니다.",
            "후보 비교표: 국내 Supplier와 중국/해외 베타 Factory를 candidate 형태로 표준화하여 제품 적합도, 공공데이터 확인도, 인증·통관 준비도, 입지·물류 적합도, 국가·거래위험을 100점 기준으로 비교합니다.",
            "RAG 무역 챗봇: 사용자의 무역·수입 질문을 구조화 데이터와 문서 근거로 보강한 뒤 LLM에 전달합니다. 답변은 조회된 evidence만 사용하고, 근거 부족 항목은 '확인 필요'로 표시합니다.",
            "Evidence Drawer: AI 답변, 후보 카드, 리스크 카드마다 providerName, datasetName, sourceType, sourceUrl, fetchedAt, license를 보여주어 심사자와 사용자가 결론의 근거를 추적할 수 있습니다.",
            "리포트 미리보기: 입력 제품 정보, 후보 비교 점수, 인증·통관 체크리스트, 국가·거래위험, AI 요약, 사용 공공데이터 출처를 보고서 형태로 제공합니다.",
        ],
    )
    add_callout(
        doc,
        "핵심 우수성",
        "지도와 챗봇을 따로 제공하는 것이 아니라, 동일한 공공데이터 evidence를 지도, 후보 비교표, 리스크 점수, AI 답변, 리포트에 재사용합니다. 따라서 사용자는 '어떤 공장이 어디에 있는지'와 '왜 그 후보를 검토해야 하는지'를 한 번에 확인할 수 있습니다.",
    )

    add_heading(doc, "2. 제품 및 서비스 제안배경(활용적정성)", 2)
    add_para(
        doc,
        "한국 중소기업과 소상공인은 특정 제품을 수입하거나 조달하려 할 때 공장 후보, 국내 생산 가능성, HS코드, 인증·통관 요건, 리콜 이력, 국가위험을 각각 다른 사이트에서 확인해야 합니다. 특히 공공데이터는 신뢰도가 높지만 데이터셋 단위로 흩어져 있고, 비전문가가 제품명 하나로 관련 공장·규제·리스크를 연결하기 어렵습니다.",
    )
    add_bullets(
        doc,
        [
            "문제 1: 공장 검색과 수입 리스크 확인이 분리되어 있어 후보 탐색 후 다시 인증·통관·리콜 데이터를 수작업으로 확인해야 합니다.",
            "문제 2: 국내 공장 데이터는 산업단지·생산품·주소 중심으로 풍부하지만, 제품명 기반 매칭과 지도 시각화가 약해 실제 조달 의사결정에 바로 쓰기 어렵습니다.",
            "문제 3: 중국 등 해외 공장 데이터는 좌표·품목 세부성·검증 수준이 제각각이므로, 국내 데이터와 같은 기준으로 비교하되 베타 데이터임을 명확히 표시해야 합니다.",
            "문제 4: 무역 초보자는 HS코드, 세관장확인대상, KC 인증, 리콜, 국가위험이 어떤 의미인지 알기 어렵고, 챗봇이 근거 없이 단정하면 오히려 위험합니다.",
        ],
    )
    add_para(
        doc,
        "따라서 SupplyMap AI는 산업통상부 및 산하 공공기관 데이터를 서비스의 신뢰 기반으로 삼고, 관세청·식약처·민간 지도 데이터를 보조로 연계해 '공장 후보 탐색 - 리스크 확인 - 근거 기반 AI 상담'을 하나의 흐름으로 제공합니다. 특히 챗봇은 단순 생성형 답변이 아니라, 플랫폼이 조회한 데이터와 출처를 먼저 묶은 뒤 LLM에 전달하는 RAG 구조이므로 사용자는 답변 옆에서 원천 근거를 확인할 수 있습니다.",
    )

    add_heading(doc, "3. 제품 및 서비스 세부내용", 2)
    add_heading(doc, "활용한 산업부 공공데이터", 3)
    add_bullets(
        doc,
        [
            "한국산업단지공단 공장등록생산정보조회서비스: 회사명, 주소, 생산품, 업종 정보를 국내 Supplier 후보로 정규화합니다.",
            "한국산업단지공단 산업동향조사 통계 조회 서비스: 산업단지별 입주기업, 가동률, 생산, 수출, 고용 정보를 IndustrialComplex 맥락으로 연결합니다.",
            "KOTRA 무역투자 질의응답 데이터셋: 무역 실무 질문에 대한 RAG 문서 근거로 사용합니다.",
            "KOTRA 해외시장뉴스 및 국가정보: 국가별 시장동향, 규제 변화, 통관 이슈, 시장 진출 정보를 챗봇과 리스크 분석에 사용합니다.",
            "한국무역보험공사 K-SURE 국가별 업종별 위험지수: 국가·거래위험과 결제위험 신호를 RiskSignal로 변환합니다.",
            "국가기술표준원·Safety Korea 제품안전 인증 및 리콜 정보: KC 인증, 국내 리콜, 국외 리콜 여부를 제품안전 리스크로 표시합니다.",
        ],
    )
    add_heading(doc, "활용한 타 기관 또는 민간 데이터", 3)
    add_bullets(
        doc,
        [
            "관세청 및 UNI-PASS 데이터: HS 후보, 세관장확인대상 법령, 확인서류, 요건승인기관, 관세율, 관세환율, 품목별 국가별 수출입실적을 분석합니다.",
            "식품의약품안전처 데이터: 수입식품 해외제조업소, 화장품 원료성분, 화장품 사용제한 원료, 의료기기 품목/허가 정보를 해외 후보와 품목 리스크 보강에 사용합니다.",
            "지도/지오코딩 데이터: Kakao Map/Local은 국내 공장 주소 좌표화와 지도 표시, Google Maps/AMap은 해외 후보 지오코딩 및 베타 지도 보강에 사용합니다.",
            "사용자 업로드 데이터: 공급업체 목록, 견적서, 카탈로그, 해외제조업소 엑셀 스냅샷을 USER_UPLOAD 출처로 분리해 후보 보강에 사용합니다.",
        ],
    )
    add_heading(doc, "구현기술 및 서비스 방법", 3)
    add_bullets(
        doc,
        [
            "Frontend: Next.js App Router, Tailwind, React 기반의 분석 워크벤치, 지도, 후보 비교표, 챗봇 UI를 구현했습니다.",
            "Data layer: Prisma + SQLite를 사용하고, SupplyDataSource, IndustrialComplex, Supplier, RiskSignal, ChatEvidence 모델로 공공데이터 출처와 evidence를 관리합니다.",
            "Adapter 구조: KICOX, KOTRA, K-SURE, Safety Korea, 관세청/UNI-PASS, 식약처 API adapter를 분리하고, API 키가 없어도 mock fallback으로 데모가 동작하게 설계했습니다.",
            "Scoring: 제품 적합도 30, 공공데이터 확인도 20, 인증·통관 준비도 20, 입지·물류 적합도 15, 국가·거래위험 15점으로 총점 100점의 의사결정 보조 점수를 계산합니다.",
            "RAG 챗봇: 사용자의 질문을 제품명·HS코드·국가·후보·리스크 evidence로 보강하고, LLM에는 조회된 근거만 전달합니다. 답변에는 evidence id와 출처 배지를 붙이고, 근거가 약하면 확인 필요로 표시합니다.",
        ],
    )
    add_heading(doc, "기존 서비스와의 차별성", 3)
    add_bullets(
        doc,
        [
            "일반 지도 서비스는 위치만 보여주지만 SupplyMap AI는 공장 후보, 생산품, 산업단지 맥락, 인증·통관·국가위험을 한 화면에 결합합니다.",
            "일반 챗봇은 근거 없이 답변할 수 있지만 SupplyMap AI 챗봇은 구조화 DB와 ChatEvidence를 먼저 조회하고 출처를 함께 보여줍니다.",
            "해외 공장 데이터를 정부 검증 데이터처럼 포장하지 않고 PRIVATE/OTHER_PUBLIC/USER_UPLOAD로 분리해 신뢰 수준을 표시합니다.",
            "심사위원과 사용자 모두 데이터 카탈로그에서 providerName, datasetName, sourceType, sourceUrl, fetchedAt, license를 확인할 수 있습니다.",
        ],
    )
    add_heading(doc, "서비스 링크 주소", 3)
    add_bullets(
        doc,
        [
            "서비스 URL: Vercel 배포 후 https://<project-name>.vercel.app/supplymap 입력",
            "GitHub URL: GitHub 업로드 후 https://github.com/<user>/<repo> 입력",
            "로컬 데모 URL: http://127.0.0.1:3001/supplymap",
        ],
    )

    add_heading(doc, "4. 아이디어의 사업화방안 및 기대효과(사업성, 실현가능성)", 2)
    add_heading(doc, "사업화 방안", 3)
    add_bullets(
        doc,
        [
            "초기 고객: 해외 수입을 처음 시도하는 중소기업, 소상공인, 온라인 셀러, 무역대행사, 국내 제조업체 영업팀을 대상으로 합니다.",
            "SaaS 모델: 월 구독형으로 제품 검색 횟수, AI 상담 횟수, 리포트 생성 횟수, 팀 계정 수에 따라 Basic/Pro/Agency 요금제를 구성합니다.",
            "B2B 확장: 무역대행사와 제조업 협회에는 후보 리스트 관리, RFQ 템플릿, 공급업체 검증 워크플로, 맞춤 데이터 커넥터를 제공합니다.",
            "데이터 확장: 현재 중국 베타 레이어를 시작으로 베트남, 일본, 미국 등 국가별 제조업소 데이터와 KOTRA/K-SURE 문서를 확장합니다.",
            "리스크 고도화: 실제 API 호출 안정화, 기업 실재성 검증, 인증서 원문 OCR, 수입요건 자동 체크리스트, PDF 리포트 자동 생성으로 발전시킵니다.",
        ],
    )
    add_heading(doc, "실현가능성", 3)
    add_bullets(
        doc,
        [
            "현재 프로젝트는 Next.js 웹앱으로 구현되어 있으며, /supplymap, /data-catalog, /api/supplymap/* 경로가 작동합니다.",
            "현재 DB에는 SupplyDataSource 17개, 국내 Supplier 3,322개, IndustrialComplex 226개, RiskSignal 2,714개, ChatEvidence 2,794개, Factory 42,396개, Product 48,064개가 들어 있습니다.",
            "API 키가 없거나 실시간 API가 실패해도 mock fallback과 seed data로 데모가 안정적으로 작동하도록 설계했습니다.",
            "Vercel 배포 또는 Render/Railway 배포가 가능하며, 장기 운영 시에는 SQLite를 PostgreSQL/Supabase/Neon/Turso 등 관리형 DB로 이전할 수 있습니다.",
        ],
    )
    add_heading(doc, "사회적 파급 및 기대효과", 3)
    add_bullets(
        doc,
        [
            "수입기업은 후보 탐색과 인증·통관 리스크 확인 시간을 줄이고, 근거 없는 해외 공급업체 선택으로 인한 비용 손실을 낮출 수 있습니다.",
            "국내 산업단지 기업은 제품명 기반 검색과 지도 노출을 통해 신규 판로를 확보할 수 있습니다.",
            "공공데이터는 단순 조회를 넘어 AI 답변, 지도 시각화, 리스크 점수, 리포트로 재사용되어 활용 가치가 커집니다.",
            "무역 초보자도 Evidence Drawer를 통해 원천 데이터와 확인 필요 항목을 함께 보며 안전한 의사결정을 할 수 있습니다.",
        ],
    )


def add_data_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "C. 활용 데이터 전체 정리", 1)
    add_callout(
        doc,
        "작성 원칙",
        "신청서의 '산업통상부 활용 데이터'에는 MOTIE_PUBLIC만 기재하고, 관세청·식약처·지도·민간·사용자 업로드 데이터는 '기타 활용 데이터'로 분리하는 것이 안전합니다.",
        fill="FFF9E8",
    )

    motie_rows = [
        ["한국산업단지공단", "공장등록생산정보조회서비스", "MOTIE_PUBLIC", "mock/connected", "국내 Supplier 후보 3,322개, 카테고리별 국내 공장 히트맵"],
        ["한국산업단지공단", "산업동향조사 통계 조회 서비스", "MOTIE_PUBLIC", "mock", "IndustrialComplex 226개, 산업단지 역량/입지 맥락"],
        ["대한무역투자진흥공사", "무역투자 질의응답 데이터셋", "MOTIE_PUBLIC", "mock", "AI 무역 챗봇 RAG 근거"],
        ["대한무역투자진흥공사", "해외시장뉴스", "MOTIE_PUBLIC", "connected", "국가·시장·통관 이슈 RiskSignal 544건"],
        ["대한무역투자진흥공사", "국가정보", "MOTIE_PUBLIC", "mock", "국가·시장·거래위험 문서 근거"],
        ["한국무역보험공사", "국가별 업종별 위험지수", "MOTIE_PUBLIC", "connected", "국가·거래위험 RiskSignal 1,972건"],
        ["국가기술표준원·Safety Korea", "제품안전 인증 및 리콜 정보", "MOTIE_PUBLIC", "connected", "KC 인증, 국내/국외 리콜, 제품안전 리스크 34건"],
    ]
    add_heading(doc, "산업통상부 및 산하 공공기관 데이터", 2)
    add_table(doc, ["기관", "데이터셋명", "sourceType", "상태", "플랫폼 활용"], motie_rows, widths_cm=[3.1, 4.2, 2.4, 2.2, 5.4], font_size=7)

    other_rows = [
        ["관세청", "HS부호 및 세관장확인대상물품", "OTHER_PUBLIC", "HS 후보, 확인법령, 요건승인기관, 확인서류"],
        ["관세청", "품목별 국가별 수출입실적", "OTHER_PUBLIC", "수입액·수입중량·무역수지 기반 시장규모/의존도"],
        ["관세청 UNI-PASS", "HS부호검색", "OTHER_PUBLIC", "HS 10단위 후보와 품명 보강"],
        ["관세청 UNI-PASS", "세관장확인대상 물품 조회", "OTHER_PUBLIC", "수입요건·확인법령 분석"],
        ["관세청 UNI-PASS", "관세율 기본 조회", "OTHER_PUBLIC", "예상 관세율·FTA 검토"],
        ["관세청 UNI-PASS", "관세환율 정보", "OTHER_PUBLIC", "수입 원가·환율 리스크 보조 분석"],
        ["관세청 UNI-PASS", "HS CODE 내비게이션", "OTHER_PUBLIC", "신고 건수 기반 HS 후보 보강"],
        ["식품의약품안전처", "수입식품 해외제조업소 정보", "OTHER_PUBLIC", "중국/해외 베타 공장 레이어 및 식품·포장 후보"],
        ["식품의약품안전처", "화장품 원료성분정보/사용제한 원료정보", "OTHER_PUBLIC", "화장품 품목 리스크와 원료 확인 보강"],
        ["식품의약품안전처", "의료기기 품목정보/품목허가정보", "OTHER_PUBLIC", "의료기기 품목 리스크 및 해외/국내 후보 보강"],
        ["Kakao Developers", "Kakao Map/Local", "PRIVATE", "국내 공장 주소 지오코딩과 Kakao 지도 표시"],
        ["Google Maps Platform", "Geocoding/Maps", "PRIVATE", "해외 후보 좌표 보강 및 지도 연동 후보"],
        ["AMap", "Web Service/Map SDK", "PRIVATE", "중국 주소 지오코딩 및 해외 후보 좌표 보강"],
        ["사용자 업로드", "해외제조업소·해외작업장·수출업소 엑셀", "USER_UPLOAD", "Factory 42,363개 스냅샷, 중국 베타 히트맵"],
        ["사용자 업로드", "공급업체 목록·견적서·카탈로그", "USER_UPLOAD", "사용자 보유 후보 보강 예정"],
        ["OpenAI/DeepSeek", "LLM API", "PRIVATE", "RAG 챗봇의 자연어 답변 생성. 근거 데이터만 전달"],
    ]
    add_heading(doc, "기타 공공데이터·민간·사용자 업로드 데이터", 2)
    add_table(doc, ["기관/원천", "데이터셋명", "sourceType", "플랫폼 활용"], other_rows, widths_cm=[3.4, 4.9, 2.4, 6.6], font_size=7)

    db_rows = [
        ["SupplyDataSource", "17", "출처 관리 기준. providerName, datasetName, sourceType, sourceUrl, fetchedAt, license 유지"],
        ["IndustrialComplex", "226", "국내 산업단지 지도/입지/산업동향 맥락"],
        ["Supplier", "3,322", "한국산업단지공단 기반 국내 공장 후보"],
        ["RiskSignal", "2,714", "인증·리콜·통관·국가·시장·거래위험 신호"],
        ["ChatEvidence", "2,794", "RAG 챗봇과 Evidence Drawer 근거"],
        ["Factory", "42,396", "중국/해외 베타 공장 레이어"],
        ["Product", "48,064", "Factory별 제품·카테고리·HS 후보"],
        ["Certificate", "44", "제품/공장 인증 스냅샷"],
        ["RiskEvent", "8", "리콜·부적합 등 레거시 위험 이벤트"],
        ["TradeRequirement", "3", "HS별 수입요건 스냅샷"],
    ]
    add_heading(doc, "현재 로컬 DB 반영 현황", 2)
    add_table(doc, ["모델/테이블", "레코드 수", "활용"], db_rows, widths_cm=[4.0, 2.5, 10.8], font_size=8)

    china_rows = [
        ["전체 중국 Factory", "42,389개", "중국/해외 베타 지도와 후보 비교"],
        ["좌표 보유 중국 Factory", "37,723개", "지도 표시 가능. 단, 일부는 도시/성 중심 추정 좌표"],
        ["주요 sourceTags", "xlsx_overseas_food_facilities_20260621: 42,363개", "사용자 업로드/식약처 해외제조업소 계열 스냅샷"],
        ["주요 제품 카테고리", "기구·용기·포장 25,609, 가공식품 11,167, 농산물 5,331, 식품첨가물 2,693, 수산물 2,481", "중국 베타 히트맵 카테고리 필터"],
    ]
    add_heading(doc, "중국/해외 베타 히트맵 데이터 현황", 2)
    add_table(doc, ["항목", "수량/내용", "의미"], china_rows, widths_cm=[4.0, 6.2, 7.1], font_size=8)


def add_chatbot_and_heatmap_explanation(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "D. 메인 기능 설명: RAG 무역 챗봇과 히트맵", 1)
    add_heading(doc, "RAG 기반 AI 무역 챗봇 원리", 2)
    add_bullets(
        doc,
        [
            "1단계 입력 이해: 사용자가 제품명, HS코드 후보, 수입 희망 국가, 공급 희망 지역, 자유 질문을 입력합니다.",
            "2단계 데이터 조회: 국내 Supplier, 중국/해외 Factory, RiskSignal, ChatEvidence, KOTRA/K-SURE/Safety Korea/관세청 근거를 검색합니다.",
            "3단계 프롬프트 강화: LLM에 원문 전체를 무작정 보내지 않고, 후보·리스크·출처·확인 필요 항목을 구조화한 evidence context만 전달합니다.",
            "4단계 안전한 답변: 시스템 프롬프트는 조회된 근거만 사용하도록 제한하고, HS코드·인증·통관·법률 판단은 단정하지 않으며 근거 부족 시 '확인 필요'로 답하게 합니다.",
            "5단계 근거 표시: 답변 옆 Evidence Drawer에서 providerName, datasetName, sourceType, sourceUrl, fetchedAt, license를 확인할 수 있습니다.",
        ],
    )
    add_callout(
        doc,
        "사용자 효과",
        "초보 수입자는 '어떤 공장을 볼지', '어떤 인증·통관 요건을 확인할지', '어떤 국가/거래 리스크가 있는지'를 한 번에 이해할 수 있습니다. 챗봇은 단순 상담원이 아니라 공공데이터를 읽어주는 무역 의사결정 코파일럿 역할을 합니다.",
    )

    add_heading(doc, "히트맵 기능 설명", 2)
    add_bullets(
        doc,
        [
            "국내 공장 히트맵: Kakao Map 기반으로 국내 Supplier 3,322개를 표시합니다. 현재 주소 좌표 3,190개, 보조/추정 좌표 132개를 구분하며 카테고리별 필터를 제공합니다.",
            "중국 공장 베타 히트맵: 중국 Factory 42,389개 중 좌표가 있는 37,723개를 기반으로 표시합니다. 성/도시 중심 추정 좌표가 포함되어 있으므로 UI에서 '베타·좌표 추정 포함'으로 명시합니다.",
            "마우스 오버/클릭: 공장명, 위치, 생산품, 주소, 좌표 신뢰도, sourceType을 보여주어 후보 검토의 첫 단서를 제공합니다.",
            "챗봇과의 연결: 지도에서 본 품목·지역 분포는 RAG 챗봇 질문으로 이어져 '이 후보의 인증·통관상 주의점', '국내 후보와 중국 후보의 차이' 같은 답변을 받을 수 있습니다.",
        ],
    )


def add_deployment_guide(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "E. GitHub 및 Vercel 제출 링크 준비 가이드", 1)
    add_callout(
        doc,
        "중요 주의",
        "API 키가 들어 있는 .env 파일은 절대 GitHub에 올리면 안 됩니다. 현재 .gitignore에는 .env와 prisma/dev.db가 제외되어 있습니다. Vercel에서 실제 데이터를 보이게 하려면 DB 배포 방식을 먼저 결정해야 합니다.",
        fill="FFF9E8",
    )

    add_heading(doc, "1. GitHub 업로드 절차", 2)
    add_numbered(
        doc,
        [
            "GitHub에서 새 repository를 생성합니다. 예: supplymap-ai 또는 portmind-supplymap-ai.",
            "로컬 프로젝트 폴더에서 npm run lint, npm run typecheck, npm run build를 실행해 오류가 없는지 확인합니다.",
            "git init을 실행하고, git status로 포함될 파일을 확인합니다. .env, node_modules, .next는 포함되면 안 됩니다.",
            "빠른 심사용 데모로 SQLite DB를 함께 올릴 경우 prisma/dev.db는 git add -f prisma/dev.db로 강제 추가해야 합니다. 단, 이 방법은 읽기 중심 데모용이며 실서비스용은 아닙니다.",
            "git add ., git commit -m \"Initial SupplyMap AI demo\", git branch -M main을 실행합니다.",
            "GitHub repository 주소를 remote로 등록합니다. git remote add origin https://github.com/<user>/<repo>.git",
            "git push -u origin main으로 업로드합니다.",
        ],
    )
    add_para(doc, "권장 사전 점검 명령:")
    add_para(
        doc,
        "npm run lint\nnpm run typecheck\nnpm run build\ngit status --short\ngit grep -n \"OPENAI_API_KEY\\|DEEPSEEK_API_KEY\\|KAKAO_REST_API_KEY\" -- ':!.env' ':!.env*.local' || true\ngit grep -n \"API_KEY\" -- ':!.env' ':!.env*.local' || true",
    )

    add_heading(doc, "2. Vercel 배포 절차", 2)
    add_numbered(
        doc,
        [
            "vercel.com에 로그인하고 Add New Project를 선택합니다.",
            "GitHub에 올린 repository를 Import합니다.",
            "Framework Preset은 Next.js로 자동 인식되게 둡니다.",
            "Build Command는 npm run build, Install Command는 기본 npm install을 사용합니다.",
            "Environment Variables에 필요한 키 이름만 등록합니다. 키 값은 문서나 GitHub에 노출하지 않습니다.",
            "Kakao Developers 콘솔에서 JavaScript SDK 도메인에 https://<project-name>.vercel.app를 추가합니다.",
            "Deploy 후 생성된 URL 뒤에 /supplymap을 붙여 서비스 링크로 제출합니다.",
        ],
    )

    add_heading(doc, "3. Vercel에서 DB가 보이게 하는 선택지", 2)
    add_bullets(
        doc,
        [
            "빠른 심사용 선택지: prisma/dev.db를 GitHub에 포함하고 DATABASE_URL=file:./dev.db로 설정합니다. Prisma schema가 prisma 폴더에 있으므로 file:./dev.db는 prisma/dev.db를 가리킵니다. DB는 읽기 중심 데모로만 사용하고, Settings처럼 쓰기 기능은 제출 링크에서 사용하지 않는 것이 안전합니다.",
            "권장 운영 선택지: Supabase/Neon/PostgreSQL 또는 Turso/libSQL 같은 관리형 DB로 이전합니다. 이 경우 Prisma datasource provider와 DATABASE_URL을 바꾸고 seed/ingest를 다시 실행해야 합니다.",
            "대안 선택지: Render/Railway처럼 persistent disk를 제공하는 서버에 배포하면 SQLite 파일을 계속 사용할 수 있습니다. 다만 심사위원 링크는 Vercel보다 설정이 길어질 수 있습니다.",
        ],
    )

    add_heading(doc, "4. Vercel 환경변수 목록", 2)
    env_rows = [
        ["DATABASE_URL", "Prisma DB 연결. 빠른 데모는 file:./dev.db, 운영은 관리형 DB URL"],
        ["NEXT_PUBLIC_KAKAO_JAVASCRIPT_KEY", "Kakao 지도 표시. 공개 키이므로 도메인 제한 필수"],
        ["KAKAO_REST_API_KEY", "서버 측 Kakao Local/geocode 호출"],
        ["DATA_GO_KR_SERVICE_KEY", "공공데이터포털 공통 serviceKey"],
        ["SAFETY_KOREA_API_KEY", "Safety Korea 인증·리콜 API"],
        ["UNIPASS_API_KEY 또는 UNIPASS_*", "UNI-PASS HS/관세율/요건 API"],
        ["MFDS_*", "식약처 식품·화장품·의료기기 API"],
        ["DEEPSEEK_API_KEY 또는 OPENAI_API_KEY", "서버 측 LLM 호출. 절대 NEXT_PUBLIC으로 만들지 않음"],
        ["SUPPLYMAP_DEEPSEEK_ENABLED", "LLM 실호출 ON/OFF. 데모 안정성이 중요하면 false 가능"],
    ]
    add_table(doc, ["환경변수명", "용도"], env_rows, widths_cm=[5.5, 11.8], font_size=8)

    add_heading(doc, "5. 기획서에 넣을 링크 문구", 2)
    add_bullets(
        doc,
        [
            "서비스 확인 링크: https://<project-name>.vercel.app/supplymap",
            "데이터 출처 페이지: https://<project-name>.vercel.app/data-catalog",
            "GitHub 저장소: https://github.com/<user>/<repo>",
            "제출 전 점검: 서비스 링크에서 국내 공장 히트맵, 중국 공장 베타 히트맵, AI 무역 코파일럿, 데이터 출처 페이지가 모두 열리는지 확인합니다.",
        ],
    )


def add_final_checklist(doc: Document) -> None:
    add_heading(doc, "F. 제출 전 체크리스트", 1)
    add_bullets(
        doc,
        [
            "공식 신청서의 제품서비스명은 'SupplyMap AI'로 통일합니다.",
            "산업통상부 활용 데이터에는 한국산업단지공단, KOTRA, K-SURE, 국가기술표준원·Safety Korea 데이터를 우선 기재합니다.",
            "기타 활용 데이터에는 관세청, UNI-PASS, 식약처, 지도 API, 사용자 업로드, LLM API를 분리해 기재합니다.",
            "서비스 링크는 반드시 /supplymap으로 끝나는 직접 링크를 넣습니다.",
            "GitHub에는 .env, API key, node_modules, .next를 올리지 않습니다.",
            "Kakao JS 도메인에 Vercel 배포 도메인을 추가합니다.",
            "챗봇 설명은 '공공데이터 기반 RAG'와 '근거 부족 시 확인 필요'를 강조합니다.",
            "히트맵 설명은 국내 Kakao 지도와 중국 베타 지도, 좌표 신뢰도 차이를 명확히 설명합니다.",
        ],
    )


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_application_form_text(doc)
    doc.add_page_break()
    add_planning_sections(doc)
    add_data_appendix(doc)
    add_chatbot_and_heatmap_explanation(doc)
    add_deployment_guide(doc)
    add_final_checklist(doc)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("PortMind SupplyMap AI - 공공데이터 기반 공급망·무역 AI 코파일럿")
        style_run(run, size=8, color=MUTED)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
